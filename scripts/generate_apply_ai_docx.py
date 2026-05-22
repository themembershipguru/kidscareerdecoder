"""DOCX: Apply AI Algorithms — milestone submission (KidsCareerDecoder)."""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def grid(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def main():
    doc = Document()
    t = doc.add_heading("Apply AI Algorithms", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph("KidsCareerDecoder — project milestone submission")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Student name: _____________________     Roll / ID: _____________________     Date: _____________________"
    )

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "KidsCareerDecoder applies AI after each completed aptitude quiz: a pluggable "
        "provider layer calls either Anthropic Claude or OpenAI with a fixed Ikigai-style "
        "system prompt, or optionally an external HTTP ML service (random_forest mode). "
        "Structured JSON (profile, three careers, top_strength, explanation) is parsed and "
        "stored with the session; deterministic fallbacks run when APIs are disabled or fail."
    )

    doc.add_heading("1. Deliverables and code map", level=1)
    grid(
        doc,
        ["Component", "Source files"],
        [
            ("Provider router + fallbacks", "backend/services/aiProfiler.js"),
            ("Runtime provider selection", "backend/services/aiProviderSettings.js (app_settings + AI_PROVIDER env)"),
            ("Claude integration", "backend/services/claudeProfiler.js"),
            ("OpenAI integration", "backend/services/openaiProfiler.js"),
            ("Session completion hook", "backend/controllers/sessionController.js (completeSession → getAptitudeProfile)"),
            ("Rule-based adaptive quiz ordering", "backend/services/adaptiveDifficulty.js (ML-ready heuristics)"),
        ],
    )

    doc.add_heading("2. Problem the algorithms solve", level=1)
    doc.add_paragraph(
        "Raw quiz output is six percentages plus age and country (see Preprocess Data for AI). "
        "The AI layer turns that vector into an interpretable strength profile and three "
        "aspirational yet realistic career suggestions with localized salary and pathway hints, "
        "plus a parent-facing explanation grounded in the score pattern."
    )

    doc.add_heading("3. Provider selection (algorithm control plane)", level=1)
    doc.add_paragraph(
        "getEffectiveAiProvider reads public.app_settings key ai_provider when present and "
        "allowed; otherwise AI_PROVIDER environment variable (default claude). Results are "
        "cached briefly (30 seconds) to limit database load. Allowed values: claude, openai, "
        "random_forest, fallback_only."
    )

    doc.add_heading("4. Primary models and APIs", level=1)
    grid(
        doc,
        ["Mode", "Model / endpoint", "Behaviour"],
        [
            (
                "claude (default path)",
                "Anthropic Messages API, model claude-sonnet-4-20250514",
                "System prompt encodes Ikigai rules; user message carries age, country, six percentages; response text stripped of markdown fences then JSON.parse.",
            ),
            (
                "openai",
                "OpenAI Chat Completions, model gpt-4o-mini, response_format json_object",
                "Same system prompt intent; native JSON mode reduces malformed replies.",
            ),
            (
                "random_forest",
                "POST {ML_SERVICE_URL}/predict (default http://localhost:5001)",
                "Swappable local or hosted classical/ML microservice returning profile, careers, top_strength, explanation.",
            ),
            (
                "fallback_only",
                "No external call",
                "buildFallbackFromScores only — for demos or policy-off regions.",
            ),
        ],
    )

    doc.add_heading("5. Prompt engineering (algorithm as policy)", level=1)
    doc.add_paragraph(
        "Both LLM backends share the same system prompt contract: profile is top two aptitudes "
        "hyphenated alphabetically; exactly three careers; country-specific realism (India vs "
        "USA salary formats); ban on trivial job titles; JSON-only response with fields "
        "profile, careers[], top_strength, explanation. This constrains the generative model "
        "so downstream parsing and UI remain stable."
    )

    doc.add_heading("6. Resilience and fallbacks", level=1)
    doc.add_paragraph(
        "If getAptitudeProfile throws (network, quota, invalid JSON), catch returns "
        "buildFallbackFromScores: sorted percentages yield a hyphenated profile label, "
        "top_strength, empty careers list, and a generic explanation; ai_provider is set to "
        "fallback. Session completion can still attach seed careers from public.careers by "
        "database top_aptitude when the careers array is empty."
    )

    doc.add_heading("7. Persistence and traceability", level=1)
    doc.add_paragraph(
        "completeSession writes ai_provider, profile, explanation, top_strength, careers, "
        "and country into quiz_sessions.metadata_json alongside scores_json and top_aptitude, "
        "so dashboards and analytics can attribute results to the algorithm path used."
    )

    doc.add_heading("8. Configuration (environment)", level=1)
    grid(
        doc,
        ["Variable", "Purpose"],
        [
            ("ANTHROPIC_API_KEY", "Authenticates Claude when AI_PROVIDER is claude."),
            ("OPENAI_API_KEY", "Authenticates OpenAI when AI_PROVIDER is openai."),
            ("AI_PROVIDER", "Default provider when app_settings does not override."),
            ("ML_SERVICE_URL", "Base URL for random_forest mode."),
        ],
    )

    doc.add_heading("9. Related algorithmic module (quiz flow)", level=1)
    doc.add_paragraph(
        "adaptiveDifficulty.js implements rule-based next-question selection from difficulty "
        "levels and response times (ability estimate). It is documented as ML-ready: a "
        "trained model could replace adjustAbility without changing the rest API contract."
    )

    doc.add_heading("10. Limitations and ethical note", level=1)
    doc.add_paragraph(
        "LLM suggestions are advisory, not predictive of success; copy is tuned for parents "
        "and must not replace professional guidance. API dependence implies cost, latency, and "
        "vendor policy constraints; fallback_only and score-only fallback support offline demos."
    )

    desktop = "/Users/shubhlata/Desktop/KidsCareerDecoder_Apply_AI_Algorithms.docx"
    repo = "/Users/shubhlata/Desktop/Projects/KidsCareerDecoder/docs/KidsCareerDecoder_Apply_AI_Algorithms.docx"
    doc.save(desktop)
    os.makedirs(os.path.dirname(repo), exist_ok=True)
    doc.save(repo)
    print(desktop)
    print(repo)


if __name__ == "__main__":
    main()
