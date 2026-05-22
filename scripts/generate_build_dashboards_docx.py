"""DOCX: Build Dashboards — milestone submission (KidsCareerDecoder)."""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def grid(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def main():
    doc = Document()
    t = doc.add_heading("Build Dashboards", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph("KidsCareerDecoder — project milestone submission")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Student name: _____________________     Roll / ID: _____________________     Date: _____________________"
    )

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "The application provides role-specific dashboards in React: parents see each child’s "
        "progress and analytics at a glance and drill into per-child reports with charts and "
        "AI careers; children get a playful quiz picker, session-based quiz flow, and "
        "celebratory results with radar visuals; administrators operate a full console for "
        "users, sessions, quizzes, insights, AI provider controls, and settings. Access is "
        "enforced with JWT-backed PrivateRoute guards and matching API authorization."
    )

    doc.add_heading("1. Technology stack", level=1)
    grid(
        doc,
        ["Layer", "Choices"],
        [
            ("UI", "React 18, React Router, Tailwind CSS utility classes"),
            ("Charts", "Recharts (RadarChart, LineChart) on parent Child Report"),
            ("HTTP", "Axios with Bearer token from localStorage (src/utils/api.js)"),
            ("Shell", "Header, Footer, AdminLayout sidebar navigation"),
        ],
    )

    doc.add_heading("2. Route map and access control", level=1)
    doc.add_paragraph(
        "PrivateRoute (src/components/PrivateRoute.jsx) redirects unauthenticated users to "
        "/login and blocks wrong roles from each area. Root path / sends users to the correct "
        "home by role."
    )
    grid(
        doc,
        ["Area", "Routes (examples)", "Roles"],
        [
            ("Parent", "/parent/dashboard, /parent/add-child, /parent/child/:id", "parent, admin"),
            ("Child", "/child/quiz, /child/quiz/:quizId, /child/results/:sessionId", "child"),
            ("Admin", "/admin, /admin/insights, /admin/users, …", "admin only"),
            ("Auth", "/login, /register, /forgot-password, /reset-password", "public"),
        ],
    )

    doc.add_heading("3. Parent dashboards", level=1)
    doc.add_paragraph(
        "ParentDashboard.jsx loads GET /auth/children and GET /analytics/children in parallel, "
        "merges rows by child_id, and shows each child’s name, age (from DOB or birth year), "
        "sign-in email, last quiz date, top aptitude, and session count with links to the "
        "detailed report and add-child flow."
    )
    doc.add_paragraph(
        "ChildReport.jsx calls GET /analytics/child/:childId and renders: summary cards, "
        "Recharts radar for latest session scores, optional line chart of top aptitude over "
        "time, session history table (dates, profile, top stripe), and CareerResultCard in "
        "parent variant for AI or seed careers from metadata_json."
    )

    doc.add_heading("4. Child experience (results as dashboard)", level=1)
    doc.add_paragraph(
        "QuizSelect lists published quizzes from GET /quiz and starts sessions via POST "
        "/session/start. TakeQuiz drives adaptive questions from session APIs. Results.jsx "
        "fetches GET /session/:sessionId, shows hero copy by top aptitude or AI profile, "
        "bar strip for six scores, radar chart, share-with-parent action, and CareerResultCard "
        "in child variant (no salary; pathway prefixed for kids)."
    )

    doc.add_heading("5. Admin console", level=1)
    grid(
        doc,
        ["Screen", "Purpose", "Typical API"],
        [
            ("AdminOverview", "Counts, recent sessions, recent users", "GET /admin/summary"),
            ("AdminInsights", "Funnel / insight metrics", "GET /admin/insights, attribution"),
            ("AdminUsers / AdminUserDetail", "Directory and edit", "GET/PATCH /admin/users…"),
            ("AdminSessions / AdminSessionDetail", "Support and audit", "GET /admin/sessions…"),
            ("AdminQuizzes / AdminQuizEditor", "CMS for quizzes, questions, options", "GET/PATCH/POST/DELETE /admin/quizzes…"),
            ("AdminApis", "AI provider status and override", "GET /admin/ai-status, PATCH settings"),
            ("AdminSettings", "Platform configuration", "admin settings endpoints"),
        ],
    )

    doc.add_heading("6. Shared components", level=1)
    doc.add_paragraph(
        "CareerResultCard.jsx supports variant=child vs parent for consistent career tiles. "
        "lib/aptitudeLabels.js and lib/quizScoring.js centralize labels and dimension order for "
        "charts and tables."
    )

    doc.add_heading("7. Data contract (dashboards ↔ backend)", level=1)
    doc.add_paragraph(
        "Dashboards consume JSON already produced by session completion: scores_json for six "
        "percentages, metadata_json for profile, explanation, careers, country, ai_provider. "
        "Analytics controllers aggregate sessions per child for parent and admin views."
    )

    doc.add_heading("8. UX and accessibility notes", level=1)
    doc.add_paragraph(
        "Loading and error states are shown on major fetches; forms use labels and semantic "
        "headings. Child UI uses high-contrast playful styling; parent and admin UIs use "
        "cleaner neutral layouts suitable for longer reading."
    )

    doc.add_heading("9. File index (primary pages)", level=1)
    grid(
        doc,
        ["File", "Role"],
        [
            ("src/pages/parent/ParentDashboard.jsx", "Parent home dashboard"),
            ("src/pages/parent/ChildReport.jsx", "Per-child analytics dashboard"),
            ("src/pages/parent/AddChild.jsx", "Onboarding form"),
            ("src/pages/child/QuizSelect.jsx", "Quiz picker"),
            ("src/pages/child/TakeQuiz.jsx", "Active quiz session UI"),
            ("src/pages/child/Results.jsx", "Post-quiz results dashboard"),
            ("src/pages/admin/AdminLayout.jsx + Admin*.jsx", "Admin console sections"),
            ("src/App.jsx", "Route table and role gates"),
        ],
    )

    desktop = "/Users/shubhlata/Desktop/KidsCareerDecoder_Build_Dashboards.docx"
    repo = "/Users/shubhlata/Desktop/Projects/KidsCareerDecoder/docs/KidsCareerDecoder_Build_Dashboards.docx"
    doc.save(desktop)
    os.makedirs(os.path.dirname(repo), exist_ok=True)
    doc.save(repo)
    print(desktop)
    print(repo)


if __name__ == "__main__":
    main()
