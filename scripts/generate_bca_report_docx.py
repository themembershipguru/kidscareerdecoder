"""Build BCA project report DOCX with CU preliminary pages + body from markdown."""
import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "docs" / "assets" / "cu-logo-online.png"
LOGO_FALLBACK = ROOT / "docs" / "assets" / "cu-logo.png"
# Sizes match college template (Project Report Template): cover ~4.77", header ~1.54"
COVER_LOGO_WIDTH = Inches(4.75)
HEADER_LOGO_WIDTH = Inches(1.54)
MD = ROOT / "docs" / "PROJECT_REPORT_KidsCareerDecoder.md"
DESKTOP = Path("/Users/shubhlata/Desktop/KidsCareerDecoder_BCA_Project_Report.docx")
REPO_OUT = ROOT / "docs" / "KidsCareerDecoder_BCA_Project_Report.docx"
BODY_TMP = ROOT / "docs" / "_report_body_tmp.docx"


def set_normal_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)


def add_centered(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def add_body_para(doc, text, justify=True):
    p = doc.add_paragraph(text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def resolve_logo_path():
    if LOGO.exists():
        return LOGO
    if LOGO_FALLBACK.exists():
        return LOGO_FALLBACK
    return None


def add_logo_to_header(header, logo_path, width=HEADER_LOGO_WIDTH):
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.add_run().add_picture(str(logo_path), width=width)


def setup_page_headers(doc, logo_path):
    """CU template: logo top-right on each page; skip first page of each section."""
    for section in doc.sections:
        section.different_first_page_header_footer = True
        add_logo_to_header(section.header, logo_path)
        first = section.first_page_header
        first.is_linked_to_previous = False
        if first.paragraphs:
            first.paragraphs[0].clear()


def cover_page(doc, logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo_path), width=COVER_LOGO_WIDTH)
    doc.add_paragraph()
    add_centered(doc, "CHANDIGARH UNIVERSITY", 16, True)
    add_centered(doc, "Mohali, Punjab, India", 12)
    doc.add_paragraph()
    add_centered(doc, "UNIVERSITY INSTITUTE OF COMPUTING", 14, True)
    add_centered(doc, "(Department name — confirm with your notice)", 11)
    doc.add_paragraph()
    add_centered(doc, "PROJECT REPORT", 18, True)
    add_centered(doc, "On", 12)
    doc.add_paragraph()
    add_centered(
        doc,
        "KidsCareerDecoder — Web-Based Child Aptitude Assessment\nand Career Insight Platform",
        14,
        True,
    )
    doc.add_paragraph()
    add_centered(doc, "Submitted in partial fulfilment of the requirements", 12)
    add_centered(doc, "for the award of the degree of", 12)
    add_centered(doc, "BACHELOR OF COMPUTER APPLICATIONS (BCA)", 13, True)
    add_centered(doc, "Final Year — Semester VI", 12)
    doc.add_paragraph()
    add_centered(doc, "Submitted by", 12, True)
    add_centered(doc, "Neel Kapoor", 12, True)
    add_centered(doc, "Roll No. / UID: O23BCA110261", 12)
    add_centered(doc, "E-mail: O23BCA110261@cuchd.in", 12)
    doc.add_paragraph()
    add_centered(doc, "Under the guidance of", 12, True)
    add_centered(doc, "Rohit Jha", 12, True)
    add_centered(doc, "Project Guide", 12)
    doc.add_paragraph()
    add_centered(doc, "Academic Session: ________________", 12)
    add_centered(doc, "22 May 2026", 12)
    doc.add_page_break()


def bonafide(doc):
    add_centered(doc, "BONAFIDE CERTIFICATE", 14, True)
    doc.add_paragraph()
    add_body_para(
        doc,
        "This is to certify that the project report entitled “KidsCareerDecoder — "
        "Web-Based Child Aptitude Assessment and Career Insight Platform” submitted by "
        "Neel Kapoor (Roll No. O23BCA110261, E-mail: O23BCA110261@cuchd.in) "
        "in partial fulfilment of the requirements for the award of the degree of "
        "Bachelor of Computer Applications (BCA), Final Year, Semester VI at "
        "Chandigarh University, is a record of bonafide work carried out by him/her "
        "under my supervision and guidance.",
    )
    add_body_para(
        doc,
        "The matter embodied in this report has not been submitted earlier for the award "
        "of any other degree or diploma to the best of my knowledge.",
    )
    doc.add_paragraph()
    doc.add_paragraph()
    add_body_para(doc, "Date: 22 May 2026")
    add_body_para(doc, "Place: Mohali")
    doc.add_paragraph()
    doc.add_paragraph()
    add_body_para(doc, "Signature of Guide: _________________________")
    add_body_para(doc, "Name of Guide: Rohit Jha")
    add_body_para(doc, "Designation: Project Guide")
    doc.add_paragraph()
    add_body_para(doc, "Head of Department: ________________________")
    add_body_para(doc, "Signature of HOD: __________________________")
    doc.add_page_break()


def declaration(doc):
    add_centered(doc, "DECLARATION", 14, True)
    doc.add_paragraph()
    add_body_para(
        doc,
        "I, Neel Kapoor (Roll No. O23BCA110261), a student of Bachelor of "
        "Computer Applications (BCA), Final Year, Semester VI, Chandigarh University, "
        "hereby declare that the project report entitled “KidsCareerDecoder — Web-Based "
        "Child Aptitude Assessment and Career Insight Platform” submitted to the "
        "University is my original work. The work has not been submitted elsewhere for "
        "any degree or diploma.",
    )
    add_body_para(
        doc,
        "I have followed the guidelines provided by the University and have duly "
        "acknowledged all sources of information. I understand that plagiarism is a "
        "serious academic offence.",
    )
    doc.add_paragraph()
    doc.add_paragraph()
    add_body_para(doc, "Date: 22 May 2026")
    add_body_para(doc, "Place: Mohali")
    doc.add_paragraph()
    add_body_para(doc, "Signature of Student: ______________________")
    add_body_para(doc, "Name: Neel Kapoor")
    add_body_para(doc, "Roll No.: O23BCA110261")
    doc.add_page_break()


def acknowledgement(doc):
    add_centered(doc, "ACKNOWLEDGEMENT", 14, True)
    doc.add_paragraph()
    add_body_para(
        doc,
        "I would like to express my sincere gratitude to Chandigarh University for "
        "providing the academic environment and resources necessary to complete this "
        "project.",
    )
    add_body_para(
        doc,
        "I am deeply thankful to my project guide, Rohit Jha, for continuous guidance, "
        "constructive feedback, and encouragement throughout the development of "
        "KidsCareerDecoder.",
    )
    add_body_para(
        doc,
        "I extend my thanks to the faculty members of the University Institute of "
        "Computing for their teaching and support during the BCA programme.",
    )
    add_body_para(
        doc,
        "I am grateful to my family and friends for their motivation and patience "
        "during this final-year project work.",
    )
    doc.add_page_break()


def apply_headers_to_document(doc_path, logo_path):
    doc = Document(doc_path)
    setup_page_headers(doc, logo_path)
    doc.save(doc_path)


def merge_docx(front_path, body_path, out_path):
    try:
        from docxcompose.composer import Composer
    except ImportError:
        subprocess.run(
            ["/usr/bin/python3", "-m", "pip", "install", "--user", "docxcompose", "-q"],
            check=False,
        )
        from docxcompose.composer import Composer

    master = Document(front_path)
    composer = Composer(master)
    composer.append(Document(body_path))
    composer.save(out_path)


def main():
    subprocess.run(
        ["/usr/bin/python3", str(ROOT / "scripts" / "create_cu_logo_placeholder.py")],
        check=True,
    )
    logo_path = resolve_logo_path()
    if not logo_path:
        raise FileNotFoundError(
            "CU logo missing. Place Project Report Template (2).docx in Downloads "
            "or add docs/assets/cu-logo-online.png"
        )

    set_normal_style(doc := Document())
    setup_page_headers(doc, logo_path)
    cover_page(doc, logo_path)
    bonafide(doc)
    declaration(doc)
    acknowledgement(doc)

    front_tmp = ROOT / "docs" / "_report_front_tmp.docx"
    doc.save(front_tmp)

    subprocess.run(
        [
            "pandoc",
            str(MD),
            "-o",
            str(BODY_TMP),
            "--standalone",
            "-f",
            "markdown+smart",
            "-t",
            "docx",
            "--resource-path",
            str(ROOT / "docs"),
        ],
        check=True,
        cwd=str(ROOT),
    )

    merge_docx(front_tmp, BODY_TMP, DESKTOP)
    merge_docx(front_tmp, BODY_TMP, REPO_OUT)

    for out in (DESKTOP, REPO_OUT):
        apply_headers_to_document(out, logo_path)

    for p in (front_tmp, BODY_TMP):
        if p.exists():
            p.unlink()

    print(DESKTOP)
    print(REPO_OUT)
    print(f"Logo (cover + page header): {logo_path}")


if __name__ == "__main__":
    main()
