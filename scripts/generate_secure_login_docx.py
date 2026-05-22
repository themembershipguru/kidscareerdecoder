"""DOCX: Implement Secure Login milestone documentation for KidsCareerDecoder."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_grid(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def main():
    doc = Document()
    h = doc.add_heading("KidsCareerDecoder — Implement Secure Login", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Technical documentation for the project milestone: secure authentication, "
        "authorization, and password recovery. Implementation aligns with the completed "
        "PostgreSQL schema (users, password_reset_tokens)."
    )

    doc.add_heading("1. Milestone objective", level=1)
    doc.add_paragraph(
        "Deliver a trustworthy login system so parents and children access only their "
        "allowed areas of the application: passwords are never stored in plaintext; "
        "sessions use signed tokens; privileged APIs verify identity and role; optional "
        "email-based password reset follows safe defaults."
    )

    doc.add_heading("2. End-to-end authentication model", level=1)
    doc.add_paragraph(
        "The React front end collects credentials and calls POST /auth/login. The Node.js "
        "Express backend validates the user against PostgreSQL (Supabase), compares the "
        "password with bcrypt, and returns a JSON Web Token (JWT). The browser stores the "
        "JWT and sends Authorization: Bearer <token> on subsequent API requests. Middleware "
        "verifies the JWT signature and expiry before handlers run; additional checks "
        "enforce parent vs child vs admin roles where needed."
    )

    doc.add_heading("3. Password storage (why bcrypt)", level=1)
    doc.add_paragraph(
        "Registration (POST /auth/register) and child creation hash passwords using bcrypt "
        "(bcryptjs) with cost factor 10. Only the hash is written to users.password_hash."
    )
    doc.add_paragraph(
        "Why: bcrypt is adaptive (slow by design), reducing brute-force risk if the "
        "database were compromised; each hash embeds a salt so identical passwords do not "
        "produce identical hashes."
    )

    doc.add_heading("4. JWT-based sessions", level=1)
    doc.add_paragraph(
        "After successful login, the server signs a JWT with HS256 using JWT_SECRET from "
        "environment variables. Payload includes id, full_name, role, and email; expiry is "
        "set to 7 days. Requests protected by verifyToken read Bearer tokens from the "
        "Authorization header."
    )
    doc.add_paragraph(
        "Why: stateless verification scales simply; the server does not maintain a session "
        "table for every browser tab. Short-lived tokens plus HTTPS in production limit "
        "replay exposure."
    )

    doc.add_heading("5. Authorization (beyond login)", level=1)
    doc.add_paragraph(
        "verifyToken attaches decoded claims to req.user. requireRole / requireAnyRole "
        "return 403 Forbidden when the role does not match—for example, only parent or "
        "admin may add children via POST /auth/add-child."
    )
    doc.add_paragraph(
        "The React PrivateRoute component redirects unauthenticated users to /login and "
        "prevents children from opening parent-only routes (and vice versa), improving UX "
        "while the API remains the source of truth for enforcement."
    )

    doc.add_heading("6. Forgot-password flow", level=1)
    doc.add_paragraph(
        "POST /auth/forgot-password accepts an email. For eligible accounts (parent/admin "
        "with a password), the server generates a cryptographically random token, stores "
        "SHA-256(token) in password_reset_tokens with an expiry (e.g. one hour), and emails "
        "a single-use link built with PUBLIC_APP_URL. The raw token is never persisted."
    )
    doc.add_paragraph(
        "Why: hashing stored reset tokens limits damage if the tokens table leaks; generic "
        "responses avoid confirming whether an email exists (account enumeration mitigation). "
        "Child internal emails (@child.kidscareerdecoder.internal) skip reset by design—"
        "parents manage those credentials."
    )
    doc.add_paragraph(
        "POST /auth/reset-password validates the token hash and expiry, bcrypt-hashes the "
        "new password (minimum length enforced), updates users.password_hash, and deletes "
        "used tokens."
    )

    doc.add_heading("7. Child accounts", level=1)
    doc.add_paragraph(
        "Children receive a unique internal sign-in email and an initial strong password "
        "generated server-side; the parent sees the initial password once when adding the "
        "child. This separates child login from shared parent email while keeping linkage "
        "via parent_user_id."
    )

    doc.add_heading("8. Client-side token handling", level=1)
    doc.add_paragraph(
        "Axios attaches the JWT from localStorage on each request; on HTTP 401 responses "
        "the client clears stored auth and redirects to login. Tokens are not embedded in "
        "URLs for normal browsing."
    )
    doc.add_paragraph(
        "Note for future hardening: localStorage is convenient but vulnerable if XSS occurs; "
        "production deployments should prioritize Content Security Policy, dependency "
        "updates, and optionally HttpOnly cookie sessions if requirements change."
    )

    doc.add_heading("9. Main API surface (auth)", level=1)
    add_grid(
        doc,
        ["Endpoint", "Protection", "Purpose"],
        [
            ("POST /auth/register", "Public", "Create parent account; bcrypt hash stored."),
            ("POST /auth/login", "Public", "Issue JWT after bcrypt.compare."),
            ("POST /auth/forgot-password", "Public", "Issue time-limited reset token email."),
            ("POST /auth/reset-password", "Public", "Set new password from token."),
            ("POST /auth/add-child", "JWT + parent/admin", "Create child user + hashed password."),
            ("GET /auth/children", "JWT + parent/admin", "List children for dashboard."),
        ],
    )

    doc.add_heading("10. Configuration dependencies", level=1)
    doc.add_paragraph(
        "JWT_SECRET — required for signing and verifying tokens.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Database connection — users and password_reset_tokens tables.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "SMTP / Brevo variables — for sending reset emails when forgot-password is used.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "PUBLIC_APP_URL — base URL for password reset links.",
        style="List Bullet",
    )

    doc.add_heading("11. Conclusion", level=1)
    doc.add_paragraph(
        "Secure login for KidsCareerDecoder combines industry-standard password hashing, "
        "signed JWT authorization, role-aware APIs, and a hashed-token password reset flow "
        "integrated with the existing relational schema—meeting the Implement Secure Login "
        "milestone while leaving room for CSP and cookie-based refinements as the project matures."
    )

    out = "/Users/shubhlata/Desktop/KidsCareerDecoder_Secure_Login_Milestone.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
