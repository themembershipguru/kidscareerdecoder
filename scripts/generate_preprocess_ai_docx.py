"""DOCX: Preprocess Data for AI — milestone submission (KidsCareerDecoder)."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def grid(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def main():
    doc = Document()
    t = doc.add_heading("Preprocess Data for AI", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph("KidsCareerDecoder — project milestone submission")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Student name: _____________________     Roll / ID: _____________________     Date: _____________________"
    )

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "This milestone covers feature engineering from quiz completion data: we aggregate "
        "child answers into six aptitude percentages, attach age and country for context, "
        "send a fixed schema to OpenAI or Claude, then normalize AI-generated careers before "
        "storing them in the database. A companion Jupyter notebook reproduces the same logic "
        "in Python for review and demonstration."
    )

    doc.add_heading("1. Deliverables", level=1)
    grid(
        doc,
        ["Deliverable", "Location"],
        [
            ("Written specification (this document)", "Submitted .docx file"),
            (
                "Runnable preprocessing demo",
                "notebooks/ai_preprocess_demo.ipynb (stdlib only; Jupyter in Cursor or JupyterLab)",
            ),
            (
                "Production implementation",
                "backend/controllers/sessionController.js; backend/services/aiProfiler.js; claudeProfiler.js / openaiProfiler.js",
            ),
        ],
    )

    doc.add_heading("2. Objective", level=1)
    doc.add_paragraph(
        "Transform raw, row-level quiz answers into a compact, deterministic representation "
        "suitable for downstream AI, and enforce a strict output contract so parent and child "
        "UIs receive consistent career objects."
    )

    doc.add_heading("3. Input data sources", level=1)
    grid(
        doc,
        ["Source", "Role in preprocessing"],
        [
            ("public.quiz_answers", "aptitude_type per answer; skipped flag excluded from counts."),
            ("public.users", "date_of_birth / birth_year for age; optional country column."),
            ("HTTP request", "Optional body.country; client may send locale-derived country hint."),
            ("IP / CDN headers", "Fallback country when body and DB lack a value (with IPinfo when configured)."),
        ],
    )

    doc.add_heading("4. Processing pipeline (production)", level=1)
    steps = [
        ("tallyAnswers", "Skip rows with skipped=true; increment counts only for known aptitude_type values."),
        ("pickTopAptitude", "Dominant stripe for headline and DB top_aptitude (deterministic tie-break)."),
        ("normalizeCountsToScores", "Each stripe count divided by number of answered items × 100; round to two decimals."),
        ("toProfilerPayload", "Rename keys to logical_pct … practical_pct for the AI layer."),
        ("Context enrichment", "computeAgeFromDob (fallback age if missing); resolveCountryForProfiler (ordered fallbacks)."),
        ("getAptitudeProfile", "Routes to Claude, OpenAI, optional ML HTTP service, or fallback_only / error fallback."),
        ("normalizeCareerEntry", "Coerce each career to title, salary, pathway, match_reason strings."),
        ("Persistence", "scores_json and metadata_json (including careers, country, ai_provider) on quiz_sessions."),
    ]
    grid(doc, ["Step / function", "Description"], steps)

    doc.add_heading("5. AI input schema", level=1)
    doc.add_paragraph(
        "The model receives a natural-language user message built from child age, country, "
        "and the six percentage values. The system prompt constrains JSON output shape "
        "(profile, careers array, top_strength, explanation)."
    )
    grid(
        doc,
        ["Conceptual field", "Meaning"],
        [
            ("logical_pct … practical_pct", "Percent of non-skipped answers attributed to each stripe."),
            ("age", "Whole-number age for tone and realism of suggestions."),
            ("country", "Guides currency and pathway examples (e.g. India vs United States)."),
        ],
    )

    doc.add_heading("6. Normalized output schema (careers)", level=1)
    grid(
        doc,
        ["Field", "Meaning"],
        [
            ("title", "Career name suitable for UI."),
            ("salary", "Human-readable range (₹ LPA or $K as appropriate)."),
            ("pathway", "One-line education or progression hint."),
            ("match_reason", "Evidence-based link to the child’s aptitude pattern."),
        ],
    )

    doc.add_heading("7. Demonstration notebook", level=1)
    doc.add_paragraph(
        "File: notebooks/ai_preprocess_demo.ipynb. It reimplements tally_answers, "
        "normalize_counts_to_scores, to_profiler_payload, and normalize_career_entry in Python "
        "so evaluators can run “Run All” and see before/after JSON without starting the Node server. "
        "No pip dependencies beyond the Python interpreter."
    )

    doc.add_heading("8. Mentor / course alignment", level=1)
    doc.add_paragraph(
        "Using OpenAI or Claude for the generative step is acceptable when feature engineering "
        "and output normalization are documented. This submission explicitly documents the input "
        "schema, transformations, and post-processing; the notebook adds reproducible clarity "
        "outside the API call."
    )

    doc.add_heading("9. Fallback and resilience", level=1)
    doc.add_paragraph(
        "If the LLM request fails, buildFallbackFromScores produces a profile label from scores; "
        "the session can still complete with seed rows from public.careers keyed by top aptitude."
    )

    doc.add_heading("10. Future extensions", level=1)
    doc.add_paragraph(
        "Local or self-hosted models would strengthen privacy and reproducibility; adaptive "
        "difficulty (backend/services/adaptiveDifficulty.js) is structured so a trained ranker "
        "could replace heuristics without changing the preprocessing contract to the career AI."
    )

    desktop = "/Users/shubhlata/Desktop/KidsCareerDecoder_Preprocess_Data_for_AI.docx"
    repo = "/Users/shubhlata/Desktop/Projects/KidsCareerDecoder/docs/KidsCareerDecoder_Preprocess_Data_for_AI.docx"
    doc.save(desktop)
    import os

    os.makedirs(os.path.dirname(repo), exist_ok=True)
    doc.save(repo)
    print(desktop)
    print(repo)


if __name__ == "__main__":
    main()
