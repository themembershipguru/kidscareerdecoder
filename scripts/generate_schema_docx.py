"""One-off generator: KidsCareerDecoder database schema DOCX for academic submission."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_grid_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val


def main():
    doc = Document()
    t = doc.add_heading("KidsCareerDecoder — Database Schema Documentation", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "PostgreSQL (Supabase) schema for a child-facing aptitude quiz web application "
        "with parent dashboards, AI-generated career suggestions, and analytics."
    )

    doc.add_heading("1. Purpose of the database", level=1)
    doc.add_paragraph(
        "The database stores users and family relationships, publishes quizzes with "
        "questions and scored options, records each child's quiz attempt as a session "
        "with answers and computed aptitude scores, and maintains seed career examples "
        "linked to aptitude types. Optional JSON columns capture marketing attribution, "
        "adaptive quiz state, and AI outputs without forcing rigid relational tables for "
        "every experimental feature."
    )

    doc.add_heading("2. Technology choices", level=1)
    doc.add_paragraph(
        "citext extension: case-insensitive unique emails.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "JSONB (scores_json, metadata_json, attribution_json): flexible structures for "
        "rapid iteration on scoring logic, AI career payloads, and campaign tracking.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Row Level Security (RLS): enabled on core tables for Supabase-aligned access "
        "patterns; the backend API typically uses a service role or controlled queries.",
        style="List Bullet",
    )

    doc.add_heading("3. Entity relationships (overview)", level=1)
    doc.add_paragraph(
        "users (parent/child) ←→ quiz_sessions ← quiz_answers → questions → question_options; "
        "quizzes group questions; careers reference aptitude_type values aligned with "
        "question_options. password_reset_tokens and app_settings are auxiliary tables."
    )

    doc.add_heading("4. Tables and rationale", level=1)

    sections = [
        (
            "public.users",
            "Stores every account: parents, children, students, instructors, and admins.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                (
                    "id",
                    "text PRIMARY KEY",
                    "Stable string identifier for URLs and JWT/subject mapping.",
                ),
                (
                    "email",
                    "citext NOT NULL UNIQUE",
                    "Login identity; citext avoids duplicate accounts differing only by case.",
                ),
                (
                    "password_hash",
                    "text",
                    "Stores bcrypt (or similar) hash; never plaintext passwords.",
                ),
                ("full_name", "text NOT NULL", "Display name on dashboards and reports."),
                (
                    "role",
                    "text CHECK",
                    "Separates parent vs child capabilities and admin tooling.",
                ),
                (
                    "parent_user_id",
                    "FK → users(id) ON DELETE SET NULL",
                    "Links child profiles to a guardian for parent dashboards.",
                ),
                (
                    "birth_year",
                    "smallint",
                    "Approximate age for UX copy and AI profiling when DOB absent.",
                ),
                (
                    "date_of_birth",
                    "date",
                    "More precise age calculation where parents provide it.",
                ),
                (
                    "country",
                    "text",
                    "Region hint for culturally relevant career examples (e.g. salary norms).",
                ),
                (
                    "attribution_json",
                    "jsonb",
                    "UTM/referrer captured at signup for growth analytics.",
                ),
                (
                    "created_at / updated_at",
                    "timestamptz",
                    "Audit and sorting of accounts.",
                ),
            ],
        ),
        (
            "public.quizzes",
            "Catalog of aptitude quizzes (themes, pacing, publication state).",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Stable quiz identifier in URLs and seeds."),
                ("slug", "text UNIQUE", "Human-readable path segment."),
                ("title / description", "text", "Shown in quiz picker and SEO."),
                (
                    "created_by_user_id",
                    "FK → users SET NULL",
                    "Optional author; NULL allows platform-owned seed content.",
                ),
                (
                    "default_difficulty",
                    "double precision",
                    "Baseline for adaptive engines.",
                ),
                (
                    "time_per_question_seconds",
                    "integer",
                    "Timer UX per quiz.",
                ),
                ("is_published", "boolean", "Draft vs live visibility."),
            ],
        ),
        (
            "public.questions",
            "Question stems belonging to one quiz.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Per-question stable ID."),
                ("quiz_id", "FK → quizzes CASCADE", "Groups questions under a quiz."),
                ("body", "text NOT NULL", "Scenario text shown to the child."),
                ("order_index", "integer", "Display order within the quiz."),
                (
                    "difficulty_level",
                    "double precision",
                    "Supports adaptive ordering (easier/harder next question).",
                ),
                ("aptitude_tag", "text", "Optional secondary tagging for analytics."),
            ],
        ),
        (
            "public.question_options",
            "Four choices per question; each maps to an aptitude dimension.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Stable option ID."),
                ("question_id", "FK → questions CASCADE", "Ownership of the option."),
                ("label", "text NOT NULL", "Answer text."),
                (
                    "aptitude_type",
                    "CHECK(logical,creative,verbal,social,scientific,practical)",
                    "Drives score aggregation into six aptitude stripes.",
                ),
                ("order_index", "integer", "Shuffle-safe display order."),
                ("is_correct", "boolean", "Reserved if factual quizzes are added."),
            ],
        ),
        (
            "public.quiz_sessions",
            "One row per attempt: links user, quiz, lifecycle, results.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Session token for APIs."),
                ("quiz_id / user_id", "FK", "Who took which quiz."),
                ("started_at / completed_at", "timestamptz", "Duration and reporting."),
                (
                    "status",
                    "in_progress | completed | abandoned",
                    "Analytics filters incomplete attempts.",
                ),
                ("top_aptitude", "text", "Denormalized headline strength."),
                (
                    "scores_json",
                    "jsonb",
                    "Per-dimension percentages or counts from answers.",
                ),
                (
                    "metadata_json",
                    "jsonb",
                    "Adaptive queue state, AI careers "
                    "(title, salary, pathway, match_reason), resolved country, etc.",
                ),
                (
                    "attribution_json",
                    "jsonb",
                    "Campaign context when the session started.",
                ),
            ],
        ),
        (
            "public.quiz_answers",
            "Granular record of each selected option or skip.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Answer row ID."),
                ("session_id", "FK → quiz_sessions CASCADE", "Groups answers."),
                ("question_id", "FK → questions CASCADE", "Which prompt."),
                (
                    "question_option_id",
                    "FK → question_options SET NULL",
                    "Chosen stripe; NULL if skipped.",
                ),
                ("aptitude_type", "text", "Snapshot for analytics without joins."),
                ("response_time_ms", "integer", "Engagement / difficulty signals."),
                ("skipped", "boolean", "Distinguishes omission from selection."),
                ("answered_at", "timestamptz", "Ordering within session."),
            ],
        ),
        (
            "public.careers",
            "Seed career titles aligned to aptitude_type for fallback UI.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Stable career seed ID."),
                ("title", "text NOT NULL", "Shown when AI metadata unavailable."),
                ("aptitude_type", "CHECK(six stripes)", "Maps careers to quiz dimensions."),
                ("sort_order", "integer", "Consistent listing order per stripe."),
            ],
        ),
        (
            "public.schema_migrations",
            "Optional local migration version log (parallel to Supabase migrations).",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("version", "integer PRIMARY KEY", "Applied schema version number."),
                ("applied_at", "timestamptz", "When recorded."),
            ],
        ),
        (
            "public.app_settings",
            "Key-value configuration (e.g. runtime AI provider override).",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("key", "text PRIMARY KEY", "Setting name."),
                ("value", "text NOT NULL", "Setting payload."),
                ("updated_at", "timestamptz", "Change tracking."),
            ],
        ),
        (
            "public.password_reset_tokens",
            "Secure forgot-password flow without exposing secrets.",
            [
                ("Column", "Type / constraint", "Why it is used"),
                ("id", "text PRIMARY KEY", "Token row ID."),
                ("user_id", "FK → users CASCADE", "Account to reset."),
                ("token_hash", "text NOT NULL", "Stores hash of emailed token only."),
                ("expires_at", "timestamptz", "Limits replay window."),
                ("created_at", "timestamptz", "Audit."),
            ],
        ),
    ]

    for title, purpose, rows in sections:
        doc.add_heading(title, level=2)
        add_para(doc, purpose, italic=True)
        headers = rows[0]
        body = rows[1:]
        add_grid_table(doc, list(headers), [list(r) for r in body])
        doc.add_paragraph()

    doc.add_heading("5. Indexes (performance rationale)", level=1)
    idx_rows = [
        (
            "idx_users_parent, idx_users_role",
            "Fast parent dashboards and role-based queries.",
        ),
        ("idx_quizzes_published", "List only live quizzes."),
        ("idx_questions_quiz, idx_options_question", "Load quiz content in order."),
        ("idx_sessions_user / quiz / status", "History and filtering by child or state."),
        (
            "idx_quiz_sessions_user_status_completed",
            "Analytics: latest completed sessions per user efficiently.",
        ),
        ("idx_answers_session / question", "Recompute scores from answers."),
        ("idx_careers_aptitude", "Fallback careers by top aptitude."),
        (
            "idx_password_reset_tokens_user / expires",
            "Invalidate tokens and clean up expired rows.",
        ),
    ]
    add_grid_table(doc, ["Index name pattern", "Why"], idx_rows)

    doc.add_heading("6. Security notes", level=1)
    doc.add_paragraph(
        "RLS is enabled on users, quizzes, questions, question_options, quiz_sessions, "
        "quiz_answers, careers, schema_migrations, and password_reset_tokens in the "
        "baseline migration. Production deployments should pair RLS policies with "
        "least-privilege API keys and validate parent-child ownership in application code."
    )

    doc.add_heading("7. Document metadata", level=1)
    doc.add_paragraph(
        "Generated from Supabase migration SQL in the KidsCareerDecoder repository. "
        "Schema evolves via dated migrations under supabase/migrations/."
    )

    out = "/Users/shubhlata/Desktop/Projects/KidsCareerDecoder/KidsCareerDecoder_DB_Schema_College_Submission.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
